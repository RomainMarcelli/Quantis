#!/usr/bin/env python3
# coding: utf-8
"""
Vyzor — générateur de rapport au format Word (.docx).

Lit le MÊME payload JSON que `financial_report.py` sur stdin et écrit un
.docx binaire sur stdout. Idéal quand le comptable ou le dirigeant veut
ouvrir le rapport dans Word pour ajouter ses commentaires / annotations.

La structure suit le PDF :
  - Couverture (titre, société, période)
  - Sommaire (Word génère sa propre TOC ; ici on liste les sections)
  - Synthèse Vyzor (score, sub-scores, résumé exécutif, KPIs, constats)
  - Bilan Actif / Passif
  - Compte de résultat
  - Analyse Création de valeur & Investissement
  - Analyse Financement & Rentabilité

Le rendu est volontairement plus sobre que le PDF (pas de gauge graphique,
pas de barres décoratives — Word n'est pas un format de mise en page riche),
mais TOUS les chiffres et textes sont là, éditables.
"""
from __future__ import annotations

import io
import json
import os
import sys
from typing import Any

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPT_DIR not in sys.path:
    sys.path.insert(0, _SCRIPT_DIR)

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ─── Palette ────────────────────────────────────────────────────────────────
GOLD = RGBColor(0xC5, 0xA0, 0x59)
INK = RGBColor(0x0F, 0x17, 0x2A)
MUTED = RGBColor(0x6B, 0x72, 0x80)
GREEN = RGBColor(0x10, 0xB9, 0x81)
RED = RGBColor(0xEF, 0x44, 0x44)
ORANGE = RGBColor(0xF9, 0x73, 0x16)


# ─── Helpers de base ────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str) -> None:
    """Applique une couleur de fond à une cellule (XML hack)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _heading(doc: Document, text: str, level: int = 1, color: RGBColor = INK) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(11)


def _section_caps(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GOLD


def _muted(doc: Document, text: str, size: float = 9) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = MUTED


def _body(doc: Document, text: str, size: float = 10, color: RGBColor = INK) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color


def _signal_color(signal: str | None) -> RGBColor:
    if signal == "positive":
        return GREEN
    if signal == "risk":
        return RED
    if signal == "warning":
        return ORANGE
    return INK


# ─── Cover ──────────────────────────────────────────────────────────────────
def build_cover(doc: Document, payload: dict[str, Any]) -> None:
    company = payload.get("companyName") or "Vyzor"
    company_info = payload.get("companyInfo") or {}

    # Brand
    p = doc.add_paragraph()
    run = p.add_run("VYZOR")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = GOLD

    p = doc.add_paragraph()
    run = p.add_run("INTELLIGENCE FINANCIÈRE POUR PME")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = GOLD

    doc.add_paragraph()  # spacer

    _heading(doc, company, level=1)

    legal_lines = []
    if company_info.get("legalForm") and company_info.get("capital"):
        legal_lines.append(f"{company_info['legalForm']} au capital de {company_info['capital']}")
    elif company_info.get("legalForm"):
        legal_lines.append(company_info["legalForm"])
    address = company_info.get("address")
    postal = company_info.get("postalCode")
    city = company_info.get("city")
    if address:
        legal_lines.append(f"Siège social : {address}")
    if postal or city:
        legal_lines.append(" ".join(filter(None, [postal, city])))
    if company_info.get("rcs"):
        legal_lines.append(company_info["rcs"])
    for line in legal_lines:
        _body(doc, line, size=10)

    doc.add_paragraph()
    _heading(doc, payload.get("reportTitle") or "Rapport d'analyse financière", level=2)
    period = payload.get("periodLabel")
    report_date = payload.get("reportDate")
    if period:
        _muted(doc, f"Période : {period}")
    if report_date:
        _muted(doc, f"Généré le {report_date}")

    doc.add_paragraph()
    doc.add_paragraph()
    _muted(doc, "Document confidentiel — Usage réservé au dirigeant et à son expert-comptable.")
    source = (payload.get("source") or {}).get("providerLabel")
    if source:
        _muted(doc, f"Source : {source} · Moteur Vyzor (audit-grade, déterministe)", size=8)

    doc.add_page_break()


# ─── TOC ────────────────────────────────────────────────────────────────────
def build_toc(doc: Document, payload: dict[str, Any]) -> None:
    _heading(doc, "Sommaire", level=1)
    doc.add_paragraph()

    items = payload.get("toc") or []
    for item in items:
        p = doc.add_paragraph()
        # numéro + titre
        n = p.add_run(f"{int(item.get('num', 0)):02d}  ")
        n.bold = True
        n.font.color.rgb = GOLD
        n.font.size = Pt(10)
        t = p.add_run(item.get("title", ""))
        t.bold = True
        t.font.size = Pt(11)
        # leader + page (à droite via tab stop)
        p.add_run("\t")
        page = p.add_run(str(item.get("page", "")))
        page.bold = True
        page.font.size = Pt(11)
        # description en dessous
        if item.get("description"):
            _muted(doc, f"   {item['description']}", size=9)

    doc.add_page_break()


# ─── Synthèse Vyzor ─────────────────────────────────────────────────────────
def build_synthese(doc: Document, payload: dict[str, Any]) -> None:
    _heading(doc, "Synthèse Vyzor", level=1)

    # Bandeau contextuel — NAF / effectif / exercice
    company_info = payload.get("companyInfo") or {}
    parts = []
    if company_info.get("nafCode"):
        if company_info.get("nafLabel"):
            parts.append(f"NAF : {company_info['nafCode']} — {company_info['nafLabel']}")
        else:
            parts.append(f"NAF : {company_info['nafCode']}")
    if company_info.get("effectif") is not None:
        e = f"Effectif : {company_info['effectif']}"
        if company_info.get("effectifBracket"):
            e += f" ({company_info['effectifBracket']})"
        parts.append(e)
    if payload.get("periodLabel"):
        parts.append(f"Exercice : {payload['periodLabel']}")
    if parts:
        _muted(doc, "  ·  ".join(parts), size=9)

    doc.add_paragraph()

    # Verdict
    score_data = payload.get("score") or {}
    verdict = score_data.get("verdict")
    if verdict:
        _body(doc, verdict, size=11)
    variation = (score_data.get("variation") or {}).get("text")
    if variation:
        _muted(doc, variation)

    doc.add_paragraph()

    # Score + sous-scores
    score = score_data.get("value")
    label = score_data.get("label", "")
    if score is not None:
        p = doc.add_paragraph()
        run = p.add_run(f"Score Vyzor : {int(round(score))}/100")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = INK
        if label:
            l = p.add_run(f"  ·  {label}")
            l.font.size = Pt(11)
            l.font.color.rgb = GOLD

    piliers = score_data.get("piliers") or []
    if piliers:
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for p in piliers:
            row = table.add_row()
            row.cells[0].text = p.get("label", "")
            row.cells[1].text = str(p.get("valueLabel") or "")
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Seuil de rentabilité
    breakeven = payload.get("breakeven") or {}
    if breakeven.get("ratio") is not None:
        doc.add_paragraph()
        _section_caps(doc, "Seuil de rentabilité")
        ratio = breakeven["ratio"]
        ratio_pct = int(round(ratio * 100))
        ca_label = breakeven.get("caLabel") or ""
        pm_label = breakeven.get("pointMortLabel") or ""
        ecart_label = breakeven.get("ecartLabel") or ""
        _body(doc, f"CA : {ca_label}  ·  Point mort : {pm_label}", size=10)
        if ratio < 1:
            _body(doc, f"{ratio_pct} % du point mort atteint — écart de {ecart_label} à combler.", size=10)
        else:
            _body(doc, f"{ratio_pct} % du point mort — point mort dépassé de {ecart_label}.", size=10, color=GREEN)

    # Résumé exécutif
    summary = payload.get("executiveSummary")
    if summary:
        doc.add_paragraph()
        _section_caps(doc, "Résumé exécutif")
        _body(doc, summary, size=10)

    # Indicateurs clés
    key_kpis = [k for k in (payload.get("keyKpis") or []) if k.get("valueLabel")]
    if key_kpis:
        doc.add_paragraph()
        _section_caps(doc, "Indicateurs clés")
        cols = 3
        rows = (len(key_kpis) + cols - 1) // cols
        table = doc.add_table(rows=rows * 2, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, kpi in enumerate(key_kpis):
            r = (i // cols) * 2
            c = i % cols
            label_cell = table.rows[r].cells[c]
            label_cell.text = (kpi.get("label") or "").upper()
            for p in label_cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.color.rgb = MUTED
                    run.bold = True
            value_cell = table.rows[r + 1].cells[c]
            value_cell.text = str(kpi.get("valueLabel") or "")
            for p in value_cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(13)
                    run.font.color.rgb = _signal_color(kpi.get("signal"))
                    run.bold = True

    # Constats
    constats = payload.get("constats") or []
    if constats:
        doc.add_paragraph()
        _section_caps(doc, "Constats")
        for c in constats:
            _add_constat(doc, c.get("message", ""), c.get("severity", "info"))

    doc.add_page_break()


def _add_constat(doc: Document, message: str, severity: str) -> None:
    p = doc.add_paragraph()
    color = _signal_color(severity if severity in ("positive", "risk") else "warning" if severity == "warning" else "positive")
    bullet = p.add_run("● ")
    bullet.font.color.rgb = color
    bullet.bold = True
    body = p.add_run(message)
    body.font.size = Pt(10)


# ─── Tableaux Bilan / CdR ───────────────────────────────────────────────────
def _financial_table(doc: Document, headers: list[str], rows: list[dict]) -> None:
    """Tableau comptable : header doré, alternance, totaux."""
    if not rows:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _set_cell_bg(cell, "C5A059")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Lignes
    for ri, row in enumerate(rows):
        kind = row.get("kind")
        is_section = kind == "section"
        is_total = kind == "total"
        is_emphasis = kind == "grand_total"
        indent = row.get("indent", 0)

        # Cellule label
        label_cell = table.rows[1 + ri].cells[0]
        label_cell.text = ("    " * indent) + (row.get("label") or "")
        for p in label_cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)
                if is_section or is_total or is_emphasis:
                    run.bold = True

        # Cellules valeurs
        values = row.get("values", [])
        for ci, v in enumerate(values, start=1):
            if ci >= len(headers):
                break
            cell = table.rows[1 + ri].cells[ci]
            cell.text = str(v or "")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    if is_total or is_emphasis:
                        run.bold = True

        if is_section:
            for cell in table.rows[1 + ri].cells:
                _set_cell_bg(cell, "F8FAFC")
        elif is_total:
            for cell in table.rows[1 + ri].cells:
                _set_cell_bg(cell, "FAF7F2")
        elif is_emphasis:
            for cell in table.rows[1 + ri].cells:
                _set_cell_bg(cell, "FAF7F2")
        elif (ri % 2) == 0:
            for cell in table.rows[1 + ri].cells:
                _set_cell_bg(cell, "FAFAFA")


def build_bilan_actif(doc: Document, payload: dict[str, Any]) -> None:
    _heading(doc, "Bilan — Actif", level=1)
    sub = "Présenté en euros"
    if payload.get("periodEndLabel"):
        sub += f" · Exercice clos le {payload['periodEndLabel']}"
    _muted(doc, sub)
    doc.add_paragraph()

    bilan = payload.get("bilanActif") or []
    rows = [{
        "label": r.get("label", ""),
        "kind": r.get("kind"),
        "indent": r.get("indent", 0),
        "values": [r.get("brut") or "", r.get("amort") or "", r.get("net") or "", r.get("netN1") or ""],
    } for r in bilan]
    _financial_table(doc, ["ACTIF", "Brut", "Amort. & Dép.", "Net", "Net N-1"], rows)
    doc.add_page_break()


def build_bilan_passif(doc: Document, payload: dict[str, Any]) -> None:
    _heading(doc, "Bilan — Passif", level=1)
    sub = "Présenté en euros"
    if payload.get("periodEndLabel"):
        sub += f" · Exercice clos le {payload['periodEndLabel']}"
    _muted(doc, sub)
    doc.add_paragraph()

    bilan = payload.get("bilanPassif") or []
    rows = [{
        "label": r.get("label", ""),
        "kind": r.get("kind"),
        "indent": r.get("indent", 0),
        "values": [r.get("net") or "", r.get("netN1") or ""],
    } for r in bilan]
    _financial_table(doc, ["PASSIF", "Net", "Net N-1"], rows)
    doc.add_page_break()


def build_compte_resultat(doc: Document, payload: dict[str, Any]) -> None:
    _heading(doc, "Compte de résultat", level=1)
    sub_parts = []
    if payload.get("periodLabel"):
        sub_parts.append(f"Période : {payload['periodLabel']}")
    sub_parts.append("Présenté en euros")
    _muted(doc, " · ".join(sub_parts))
    doc.add_paragraph()

    cdr = payload.get("compteResultat") or []
    rows = [{
        "label": r.get("label", ""),
        "kind": r.get("kind"),
        "indent": r.get("indent", 0),
        "values": [r.get("montant") or "", r.get("pctCa") or ""],
    } for r in cdr]
    _financial_table(doc, ["", "Montant", "% CA"], rows)
    doc.add_page_break()


# ─── KPI tiles (sections analyse) ───────────────────────────────────────────
def _kpi_tiles_section(doc: Document, title: str, items: list[dict]) -> None:
    items = [it for it in items if it.get("valueLabel")]
    if not items:
        return
    _heading(doc, title, level=2, color=GOLD)
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i in range(0, len(items), 2):
        row = table.add_row()
        for j in (0, 1):
            if i + j >= len(items):
                continue
            it = items[i + j]
            cell = row.cells[j]
            cell.text = ""
            label_p = cell.paragraphs[0]
            label_run = label_p.add_run((it.get("label") or "").upper())
            label_run.bold = True
            label_run.font.size = Pt(8)
            label_run.font.color.rgb = MUTED
            value_p = cell.add_paragraph()
            value_run = value_p.add_run(it.get("valueLabel") or "")
            value_run.bold = True
            value_run.font.size = Pt(13)
            value_run.font.color.rgb = _signal_color(it.get("signal"))
            if it.get("description"):
                desc_p = cell.add_paragraph()
                desc_run = desc_p.add_run(it["description"])
                desc_run.font.size = Pt(8.5)
                desc_run.font.color.rgb = MUTED


def build_analyse_value(doc: Document, payload: dict[str, Any]) -> None:
    rendered = False
    items_v = payload.get("valueCreationItems") or []
    if any(it.get("valueLabel") for it in items_v):
        _kpi_tiles_section(doc, "Création de valeur & Rentabilité opérationnelle", items_v)
        rendered = True
    items_i = payload.get("investmentItems") or []
    if any(it.get("valueLabel") for it in items_i):
        if rendered:
            doc.add_paragraph()
        _kpi_tiles_section(doc, "Investissement & Gestion du Besoin en Fonds de Roulement", items_i)
        rendered = True
    constats = payload.get("constatsValueCreation") or []
    if constats:
        if rendered:
            doc.add_paragraph()
        _section_caps(doc, "Constats")
        for c in constats:
            _add_constat(doc, c.get("message", ""), c.get("severity", "info"))
        rendered = True
    if rendered:
        doc.add_page_break()


def build_analyse_finance(doc: Document, payload: dict[str, Any]) -> None:
    rendered = False
    items_f = payload.get("financingItems") or []
    if any(it.get("valueLabel") for it in items_f):
        _kpi_tiles_section(doc, "Financement", items_f)
        rendered = True
    items_p = payload.get("profitabilityItems") or []
    if any(it.get("valueLabel") for it in items_p):
        if rendered:
            doc.add_paragraph()
        _kpi_tiles_section(doc, "Rentabilité", items_p)
        rendered = True
    constats = payload.get("constatsFinancing") or []
    if constats:
        if rendered:
            doc.add_paragraph()
        _section_caps(doc, "Constats")
        for c in constats:
            _add_constat(doc, c.get("message", ""), c.get("severity", "info"))


# ─── Dashboard mode ─────────────────────────────────────────────────────────
def build_dashboard_section(doc: Document, section: dict[str, Any]) -> None:
    title = section.get("title") or "Tableau de bord"
    description = section.get("description")
    _heading(doc, title, level=1)
    if description:
        _muted(doc, description)
    doc.add_paragraph()

    widgets = section.get("widgets") or []
    if not widgets:
        _muted(doc, "Ce tableau de bord ne contient aucun widget exportable.")
        return

    # Pour Word, on rend chaque widget sous forme de bloc heading + valeur + description.
    for w in widgets:
        viz = w.get("vizType")
        label = w.get("label") or ""
        if viz == "kpiCard":
            _section_caps(doc, label)
            value = w.get("valueLabel") or ""
            value_run = doc.add_paragraph().add_run(value)
            value_run.bold = True
            value_run.font.size = Pt(15)
            value_run.font.color.rgb = INK
            if w.get("description"):
                _muted(doc, w["description"])
        elif viz == "quantisScore":
            _section_caps(doc, label)
            score_text = f"{int(round(w['score']))}/100" if w.get("score") is not None else "N/D"
            run = doc.add_paragraph().add_run(score_text)
            run.bold = True
            run.font.size = Pt(15)
            if w.get("scoreLabel"):
                _muted(doc, w["scoreLabel"])
        elif viz in ("aiInsight", "alertList", "actionList"):
            _section_caps(doc, label)
            for it in (w.get("items") or []):
                _add_constat(doc, it, "info")
        else:
            # charts → placeholder textuel
            _section_caps(doc, label)
            _muted(doc, w.get("placeholderNote") or f"Visualisation {viz} disponible dans l'application.")
        doc.add_paragraph()


# ─── Build pipeline ─────────────────────────────────────────────────────────
def _build_synthese(doc: Document, payload: dict) -> None:
    build_cover(doc, payload)
    build_toc(doc, payload)
    build_synthese(doc, payload)
    build_bilan_actif(doc, payload)
    build_bilan_passif(doc, payload)
    build_compte_resultat(doc, payload)
    build_analyse_value(doc, payload)
    build_analyse_finance(doc, payload)


def _build_dashboard(doc: Document, payload: dict) -> None:
    build_cover(doc, payload)
    build_toc(doc, payload)
    sections = payload.get("dashboards") or []
    for i, section in enumerate(sections):
        if i > 0:
            doc.add_page_break()
        build_dashboard_section(doc, section)


def _build_statement(doc: Document, payload: dict) -> None:
    """Mode statement : cover + sommaire + bilan (actif + passif) OU CDR seul.
    Pas de synthèse ni d'analyse — export ciblé sur l'état financier consulté."""
    build_cover(doc, payload)
    build_toc(doc, payload)
    kind = payload.get("statementKind") or "bilan"
    if kind == "bilan":
        build_bilan_actif(doc, payload)
        build_bilan_passif(doc, payload)
    else:
        build_compte_resultat(doc, payload)


def build_document(payload: dict[str, Any]) -> bytes:
    doc = Document()
    # Marges raisonnables
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    mode = payload.get("mode") or "synthese"
    if mode == "dashboard":
        _build_dashboard(doc, payload)
    elif mode == "statement":
        _build_statement(doc, payload)
    else:
        _build_synthese(doc, payload)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main() -> int:
    raw = sys.stdin.buffer.read()
    if not raw:
        sys.stderr.write("Empty stdin\n")
        return 2
    try:
        payload = json.loads(raw)
    except json.JSONDecodeError as exc:
        sys.stderr.write(f"Invalid JSON: {exc}\n")
        return 2
    try:
        out = build_document(payload)
    except Exception as exc:
        import traceback
        sys.stderr.write(f"DOCX build failed: {type(exc).__name__}: {exc}\n")
        traceback.print_exc(file=sys.stderr)
        return 3
    sys.stdout.buffer.write(out)
    sys.stdout.buffer.flush()
    return 0


if __name__ == "__main__":
    sys.exit(main())
